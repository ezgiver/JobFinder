"""CV text extraction from PDF and DOCX files."""

import io

import fitz
from docx import Document


def extract_cv_text_from_path(path: str) -> str:
    """Read a PDF or DOCX file from disk and return extracted text."""
    import os

    name = os.path.basename(path).lower()
    with open(path, "rb") as f:
        data = f.read()

    if name.endswith(".pdf"):
        pages = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                pages.append(page.get_text())
        return "".join(pages)

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    return ""


def extract_cv_text(data: bytes, filename: str) -> str:
    """Extract plain text from raw bytes of a PDF or DOCX file."""
    name = filename.lower()

    if name.endswith(".pdf"):
        pages = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                pages.append(page.get_text())
        return "".join(pages)

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    return ""
